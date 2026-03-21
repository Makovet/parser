from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from qms_doc_parser.models.parser_models import BlockType, ParserBlock, ParserDocument, RunFormattingSnapshot


REVIEW_HEADER_STYLE = "ReviewBlockHeader"
REVIEW_WARNING_STYLE = "ReviewWarning"
REVIEW_TEXT_STYLE = "ReviewText"
REVIEW_HEADING_STYLE = "ReviewHeading"
REVIEW_CAPTION_STYLE = "ReviewCaption"
REVIEW_TABLE_META_STYLE = "ReviewTableMeta"

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_VERTICAL_ALIGNMENT_MAP = {
    "top": WD_ALIGN_VERTICAL.TOP,
    "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}


def export_review_docx(document: ParserDocument, output_path: str | Path) -> None:
    review_doc = Document()
    _configure_review_styles(review_doc)
    _add_document_summary(review_doc, document)

    for block in document.blocks:
        _render_block(review_doc, block)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    review_doc.save(output)


def export_review_docx_from_json(input_path: str | Path, output_path: str | Path) -> None:
    payload = Path(input_path).read_text(encoding="utf-8")
    document = ParserDocument.model_validate(json.loads(payload))
    export_review_docx(document=document, output_path=output_path)


def _configure_review_styles(document: DocumentObject) -> None:
    _ensure_paragraph_style(document, REVIEW_HEADER_STYLE, base_style="No Spacing", font_name="Consolas", font_size=9, bold=True, color="444444")
    _ensure_paragraph_style(document, REVIEW_WARNING_STYLE, base_style="No Spacing", font_name="Arial", font_size=10, bold=True, color="9C0006")
    _ensure_paragraph_style(document, REVIEW_TEXT_STYLE, base_style="Normal", font_name="Arial", font_size=11)
    _ensure_paragraph_style(document, REVIEW_HEADING_STYLE, base_style="Heading 2", font_name="Arial", font_size=13, bold=True)
    _ensure_paragraph_style(document, REVIEW_CAPTION_STYLE, base_style="No Spacing", font_name="Arial", font_size=10, italic=True, color="404040")
    _ensure_paragraph_style(document, REVIEW_TABLE_META_STYLE, base_style="No Spacing", font_name="Consolas", font_size=9, color="666666")


def _ensure_paragraph_style(
    document: DocumentObject,
    style_name: str,
    *,
    base_style: str,
    font_name: str,
    font_size: int,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    styles = document.styles
    if style_name in styles:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base_style]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.italic = italic
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)


def _add_document_summary(document: DocumentObject, parsed: ParserDocument) -> None:
    title = document.add_paragraph(style=REVIEW_HEADING_STYLE)
    title.add_run(f"Review DOCX export for {parsed.source.file_name}")

    summary = document.add_paragraph(style=REVIEW_TABLE_META_STYLE)
    summary.add_run(
        f"document_id={parsed.document_id} | template={parsed.template_id} | blocks={len(parsed.blocks)} | "
        f"source={parsed.source.file_name}"
    )


def _render_block(document: DocumentObject, block: ParserBlock) -> None:
    document.add_paragraph(_format_block_header(block), style=REVIEW_HEADER_STYLE)
    _render_review_hints(document, block)

    if block.block_type == BlockType.table:
        _render_table_block(document, block)
    elif block.block_type in {BlockType.empty, BlockType.unknown}:
        _render_diagnostic_block(document, block)
    else:
        _render_textual_block(document, block)


def _format_block_header(block: ParserBlock) -> str:
    zone = block.document_zone.value if hasattr(block.document_zone, "value") else str(block.document_zone)
    section_context = _safe_attr(block, "section_context")
    section_path = _safe_attr(section_context, "section_path") or []
    source_style = block.source_style or "-"
    return (
        f"[BLOCK {block.block_order:04d}] id={block.block_id} | type={block.block_type.value} | "
        f"zone={zone} | section={' > '.join(section_path) if section_path else '-'} | style={source_style}"
    )


def _render_review_hints(document: DocumentObject, block: ParserBlock) -> None:
    hints_model = _safe_attr(block, "review_render_hints")
    flags_model = _safe_attr(block, "flags")

    hints: list[str] = []
    if bool(_safe_attr(hints_model, "needs_review", False)):
        hints.append("needs_review")
    if bool(_safe_attr(hints_model, "is_suspicious", False)):
        hints.append("suspicious")
    if bool(_safe_attr(hints_model, "is_unresolved", False)):
        hints.append("unresolved")
    if bool(_safe_attr(flags_model, "is_empty", False)) and block.block_type != BlockType.empty:
        hints.append("flag:is_empty")

    if hints:
        document.add_paragraph(f"[REVIEW] {' | '.join(hints)}", style=REVIEW_WARNING_STYLE)


def _render_textual_block(document: DocumentObject, block: ParserBlock) -> None:
    paragraph_style = _select_text_style(block)
    paragraph = document.add_paragraph(style=paragraph_style)
    _apply_paragraph_formatting(paragraph, block)

    prefix = _list_prefix(block)
    if prefix:
        paragraph.add_run(prefix)

    if block.runs:
        for run_snapshot in block.runs:
            _append_run(paragraph, run_snapshot)
    else:
        text = block.normalized_text or block.raw_text or "[NO TEXT]"
        paragraph.add_run(text)

    if not paragraph.text.strip():
        paragraph.add_run("[NO TEXT]")


def _render_diagnostic_block(document: DocumentObject, block: ParserBlock) -> None:
    diagnostic_label = "[EMPTY BLOCK]" if block.block_type == BlockType.empty else "[UNKNOWN BLOCK]"
    diagnostic = document.add_paragraph(style=REVIEW_WARNING_STYLE)
    diagnostic.add_run(diagnostic_label)
    text = block.normalized_text or block.raw_text
    if text:
        diagnostic.add_run(f" {text}")


def _render_table_block(document: DocumentObject, block: ParserBlock) -> None:
    table_info = block.table_info
    if table_info is None:
        document.add_paragraph("[TABLE BLOCK WITHOUT TABLE_INFO]", style=REVIEW_WARNING_STYLE)
        return

    table_meta = document.add_paragraph(style=REVIEW_TABLE_META_STYLE)
    table_meta.add_run(
        f"table_id={table_info.table_id or '-'} | rows={table_info.rows_count or 0} | cols={table_info.cols_count or 0} | "
        f"style={table_info.table_style or '-'} | header_row={bool(table_info.has_header_row)}"
    )

    rows = max(table_info.rows_count or 0, len(table_info.cells_normalized))
    cols = max(table_info.cols_count or 0, max((len(row) for row in table_info.cells_normalized), default=0))
    if rows == 0 or cols == 0:
        document.add_paragraph("[EMPTY TABLE]", style=REVIEW_WARNING_STYLE)
        return

    table = document.add_table(rows=rows, cols=cols)
    _apply_table_style(table, table_info.table_style)

    for row_index in range(rows):
        normalized_row = table_info.cells_normalized[row_index] if row_index < len(table_info.cells_normalized) else []
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            text = normalized_row[col_index] if col_index < len(normalized_row) else None
            cell.text = text or ""
            raw_cell = _find_raw_cell(block, row_index, col_index)
            if raw_cell and raw_cell.formatting:
                if raw_cell.formatting.horizontal_alignment in _ALIGNMENT_MAP and cell.paragraphs:
                    cell.paragraphs[0].alignment = _ALIGNMENT_MAP[raw_cell.formatting.horizontal_alignment]
                if raw_cell.formatting.vertical_alignment in _VERTICAL_ALIGNMENT_MAP:
                    cell.vertical_alignment = _VERTICAL_ALIGNMENT_MAP[raw_cell.formatting.vertical_alignment]

    document.add_paragraph("", style=REVIEW_TEXT_STYLE)


def _apply_table_style(table, table_style: str | None) -> None:
    for candidate in [table_style, "Table Grid"]:
        if not candidate:
            continue
        try:
            table.style = candidate
            return
        except KeyError:
            continue


def _find_raw_cell(block: ParserBlock, row_index: int, col_index: int):
    table_info = block.table_info
    if table_info is None:
        return None
    for raw_row in table_info.cells_raw:
        for raw_cell in raw_row:
            if raw_cell.row_index == row_index and raw_cell.col_index == col_index:
                return raw_cell
    return None


def _select_text_style(block: ParserBlock) -> str:
    if block.block_type in {BlockType.heading, BlockType.appendix_heading}:
        return REVIEW_HEADING_STYLE
    if block.block_type in {BlockType.table_caption, BlockType.figure_caption, BlockType.note_label}:
        return REVIEW_CAPTION_STYLE
    return REVIEW_TEXT_STYLE


def _list_prefix(block: ParserBlock) -> str:
    if block.block_type != BlockType.list_item:
        return ""

    marker = None
    level = 0
    if block.list_formatting is not None:
        marker = block.list_formatting.marker_text
        level = block.list_formatting.level or 0
    elif block.list_info is not None:
        marker = block.list_info.list_marker
        level = block.list_info.list_level or 0

    marker = marker or "-"
    return f"{'  ' * level}{marker} "


def _apply_paragraph_formatting(paragraph, block: ParserBlock) -> None:
    formatting = block.paragraph_formatting
    if formatting is None:
        return
    if formatting.alignment in _ALIGNMENT_MAP:
        paragraph.alignment = _ALIGNMENT_MAP[formatting.alignment]
    if formatting.left_indent_pt is not None:
        paragraph.paragraph_format.left_indent = Pt(formatting.left_indent_pt)
    if formatting.first_line_indent_pt is not None and block.block_type != BlockType.list_item:
        paragraph.paragraph_format.first_line_indent = Pt(formatting.first_line_indent_pt)
    if formatting.space_before_pt is not None:
        paragraph.paragraph_format.space_before = Pt(formatting.space_before_pt)
    if formatting.space_after_pt is not None:
        paragraph.paragraph_format.space_after = Pt(formatting.space_after_pt)
    if formatting.line_spacing is not None:
        paragraph.paragraph_format.line_spacing = formatting.line_spacing


def _append_run(paragraph, run_snapshot: RunFormattingSnapshot) -> None:
    run = paragraph.add_run(run_snapshot.text)
    if run_snapshot.bold is not None:
        run.bold = run_snapshot.bold
    if run_snapshot.italic is not None:
        run.italic = run_snapshot.italic
    if run_snapshot.underline is not None:
        run.underline = run_snapshot.underline
    if run_snapshot.font_name:
        run.font.name = run_snapshot.font_name
    if run_snapshot.font_size_pt is not None:
        run.font.size = Pt(run_snapshot.font_size_pt)
    if run_snapshot.color_rgb:
        try:
            run.font.color.rgb = RGBColor.from_string(run_snapshot.color_rgb)
        except ValueError:
            pass


def _safe_attr(value: Any, attr_name: str, default: Any = None) -> Any:
    if value is None:
        return default
    return getattr(value, attr_name, default)


__all__ = ["export_review_docx", "export_review_docx_from_json"]
