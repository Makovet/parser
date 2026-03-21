from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from datetime import datetime
from pathlib import Path

from docx import Document

from qms_doc_parser.exporters.review_docx import export_review_docx, export_review_docx_from_json
from qms_doc_parser.models.parser_models import (
    BlockFlags,
    BlockType,
    CellFormattingSnapshot,
    DocumentMetadata,
    DocumentZone,
    ParserBlock,
    ParserDocument,
    ParagraphFormattingSnapshot,
    ReviewRenderHints,
    RunFormattingSnapshot,
    SectionContext,
    SourceMeta,
    StructureSummary,
    TableCellRaw,
    TableInfo,
)


class ReviewDocxExportTests(unittest.TestCase):
    def test_exporter_creates_docx_file(self) -> None:
        document = _build_parser_document()

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "review.docx"
            export_review_docx(document, output_path)

            self.assertTrue(output_path.exists())
            self.assertGreater(output_path.stat().st_size, 0)

    def test_review_docx_contains_block_headers_and_hints(self) -> None:
        document = _build_parser_document()

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "review.docx"
            export_review_docx(document, output_path)
            review_doc = Document(output_path)

            paragraphs = [paragraph.text for paragraph in review_doc.paragraphs]
            self.assertTrue(any("[BLOCK 0001] id=b1 | type=heading" in text for text in paragraphs))
            self.assertTrue(any("[REVIEW] needs_review | suspicious | unresolved" in text for text in paragraphs))
            self.assertTrue(any("[EMPTY BLOCK]" in text for text in paragraphs))
            self.assertTrue(any("[UNKNOWN BLOCK]" in text for text in paragraphs))

    def test_review_docx_preserves_runs_and_list_items_without_failure(self) -> None:
        document = _build_parser_document()

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "review.docx"
            export_review_docx(document, output_path)
            review_doc = Document(output_path)

            texts = [paragraph.text for paragraph in review_doc.paragraphs]
            self.assertTrue(any("Жирный текст" in text for text in texts))
            self.assertTrue(any("1. Первый элемент" in text for text in texts))

    def test_review_docx_renders_table_blocks(self) -> None:
        document = _build_parser_document()

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "review.docx"
            export_review_docx(document, output_path)
            review_doc = Document(output_path)

            self.assertEqual(len(review_doc.tables), 1)
            table = review_doc.tables[0]
            self.assertEqual(table.cell(0, 0).text, "Колонка")
            self.assertEqual(table.cell(1, 1).text, "B1")

    def test_handles_missing_section_context_without_failure(self) -> None:
        document = _build_parser_document()
        document.blocks[1].section_context = None

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "missing-section.docx"
            export_review_docx(document, output_path)
            review_doc = Document(output_path)

            paragraphs = [paragraph.text for paragraph in review_doc.paragraphs]
            self.assertTrue(any("id=b2" in text and "section=-" in text for text in paragraphs))

    def test_handles_missing_review_render_hints_and_flags_without_failure(self) -> None:
        document = _build_parser_document()
        document.blocks[1].review_render_hints = None
        document.blocks[1].flags = None

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "missing-review-hints.docx"
            export_review_docx(document, output_path)
            review_doc = Document(output_path)

            paragraphs = [paragraph.text for paragraph in review_doc.paragraphs]
            self.assertTrue(any("id=b2" in text for text in paragraphs))

    def test_export_from_minimal_json_helper_smoke(self) -> None:
        minimal_json = {
            "document_id": "minimal-review",
            "source": {
                "file_name": "minimal.docx",
                "parser_version": "test",
                "processed_at": "2026-03-21T00:00:00Z"
            },
            "blocks": [
                {
                    "block_id": "b1",
                    "block_order": 1,
                    "document_zone": "main_body",
                    "block_type": "paragraph",
                    "normalized_text": "Минимальный блок"
                }
            ]
        }

        with tempfile.TemporaryDirectory() as tmp_dir:
            json_path = Path(tmp_dir) / "minimal.json"
            output_path = Path(tmp_dir) / "minimal.docx"
            json_path.write_text(json.dumps(minimal_json, ensure_ascii=False, indent=2), encoding="utf-8")

            export_review_docx_from_json(json_path, output_path)
            review_doc = Document(output_path)

            paragraphs = [paragraph.text for paragraph in review_doc.paragraphs]
            self.assertTrue(any("Минимальный блок" in text for text in paragraphs))

    def test_export_from_json_helper_and_cli_work(self) -> None:
        document = _build_parser_document()

        with tempfile.TemporaryDirectory() as tmp_dir:
            json_path = Path(tmp_dir) / "document.json"
            output_helper = Path(tmp_dir) / "helper.docx"
            output_cli = Path(tmp_dir) / "cli.docx"
            json_path.write_text(document.model_dump_json(indent=2), encoding="utf-8")

            export_review_docx_from_json(json_path, output_helper)
            self.assertTrue(output_helper.exists())

            result = subprocess.run(
                [
                    sys.executable,
                    "scripts/run_export_review_docx.py",
                    str(json_path),
                    str(output_cli),
                ],
                check=False,
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, msg=result.stderr)
            self.assertTrue(output_cli.exists())


def _build_parser_document() -> ParserDocument:
    blocks = [
        ParserBlock(
            block_id="b1",
            block_order=1,
            document_zone=DocumentZone.main_body,
            block_type=BlockType.heading,
            source_style="Heading 1",
            normalized_text="Раздел 1",
            section_context=SectionContext(section_path=["1", "1.1"]),
            review_render_hints=ReviewRenderHints(needs_review=True, is_suspicious=True, is_unresolved=True),
            flags=BlockFlags(needs_review=True, is_suspicious=True),
            runs=[RunFormattingSnapshot(text="Раздел 1", bold=True)],
        ),
        ParserBlock(
            block_id="b2",
            block_order=2,
            document_zone=DocumentZone.main_body,
            block_type=BlockType.paragraph,
            source_style="Normal",
            normalized_text="Жирный текст",
            paragraph_formatting=ParagraphFormattingSnapshot(alignment="justify", left_indent_pt=12.0),
            runs=[
                RunFormattingSnapshot(text="Жирный", bold=True),
                RunFormattingSnapshot(text=" текст"),
            ],
        ),
        ParserBlock(
            block_id="b3",
            block_order=3,
            document_zone=DocumentZone.main_body,
            block_type=BlockType.list_item,
            source_style="Normal",
            normalized_text="Первый элемент",
            paragraph_formatting=ParagraphFormattingSnapshot(left_indent_pt=18.0),
            list_formatting={"level": 0, "marker_text": "1."},
        ),
        ParserBlock(
            block_id="b4",
            block_order=4,
            document_zone=DocumentZone.main_body,
            block_type=BlockType.table,
            table_info=TableInfo(
                table_id="tbl_0001",
                table_index=1,
                rows_count=2,
                cols_count=2,
                table_style="Table Grid",
                has_header_row=True,
                cells_normalized=[["Колонка", "Значение"], ["A1", "B1"]],
                cells_raw=[
                    [
                        TableCellRaw(
                            text="Колонка",
                            row_index=0,
                            col_index=0,
                            formatting=CellFormattingSnapshot(horizontal_alignment="center", vertical_alignment="center"),
                        ),
                        TableCellRaw(text="Значение", row_index=0, col_index=1),
                    ],
                    [
                        TableCellRaw(text="A1", row_index=1, col_index=0),
                        TableCellRaw(text="B1", row_index=1, col_index=1),
                    ],
                ],
            ),
        ),
        ParserBlock(
            block_id="b5",
            block_order=5,
            document_zone=DocumentZone.main_body,
            block_type=BlockType.empty,
            flags=BlockFlags(is_empty=True),
        ),
        ParserBlock(
            block_id="b6",
            block_order=6,
            document_zone=DocumentZone.unknown_zone,
            block_type=BlockType.unknown,
            raw_text="Неизвестный фрагмент",
        ),
    ]

    return ParserDocument(
        document_id="review-sample",
        source=SourceMeta(
            file_name="review-sample.docx",
            parser_version="test",
            processed_at=datetime.utcnow(),
        ),
        document_metadata=DocumentMetadata(title="Review sample"),
        structure_summary=StructureSummary(total_blocks=len(blocks), total_tables=1, total_list_items=1, total_sections=1),
        blocks=blocks,
    )


if __name__ == "__main__":
    unittest.main()
