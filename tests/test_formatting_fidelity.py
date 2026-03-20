from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from qms_doc_parser.models.parser_models import BlockType
from qms_doc_parser.pipeline.parser_pipeline import parse_docx_to_document


REGISTRY_PATH = Path("configs/style_registry_adm_tem_011_b.yaml")


class FormattingFidelityTests(unittest.TestCase):
    def _parse_temp_doc(self, builder) -> object:
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "formatting.docx"
            document = Document()
            builder(document)
            document.save(path)
            return parse_docx_to_document(input_path=path, registry_path=REGISTRY_PATH)

    def test_paragraph_formatting_snapshot_present_for_paragraph_block(self) -> None:
        parsed = self._parse_temp_doc(self._build_paragraph_doc)

        paragraph_block = next(block for block in parsed.blocks if block.block_type == BlockType.paragraph)
        self.assertIsNotNone(paragraph_block.paragraph_formatting)
        self.assertEqual(paragraph_block.paragraph_formatting.alignment, "center")
        self.assertEqual(paragraph_block.paragraph_formatting.left_indent_pt, 18.0)
        self.assertEqual(paragraph_block.paragraph_formatting.first_line_indent_pt, 9.0)

    def test_run_snapshots_are_preserved_and_serializable(self) -> None:
        parsed = self._parse_temp_doc(self._build_run_doc)

        paragraph_block = next(block for block in parsed.blocks if block.block_type == BlockType.paragraph)
        self.assertGreaterEqual(len(paragraph_block.runs), 2)
        first_run = paragraph_block.runs[0]
        self.assertEqual(first_run.text, "Важный")
        self.assertTrue(first_run.bold)
        self.assertEqual(first_run.font_name, "Arial")
        payload = parsed.model_dump(mode="json")
        self.assertIn("runs", payload["blocks"][0])

    def test_list_item_block_gets_list_formatting_snapshot(self) -> None:
        parsed = parse_docx_to_document(input_path=Path("data/input/1.docx"), registry_path=REGISTRY_PATH)

        list_block = next(block for block in parsed.blocks if block.block_type == BlockType.list_item)
        self.assertIsNotNone(list_block.list_formatting)
        self.assertIsNotNone(list_block.list_formatting.marker_text)

    def test_table_formatting_snapshot_present_on_basic_level(self) -> None:
        parsed = self._parse_temp_doc(self._build_table_doc)

        table_block = next(block for block in parsed.blocks if block.block_type == BlockType.table)
        self.assertIsNotNone(table_block.table_info)
        self.assertEqual(table_block.table_info.table_style, "Table Grid")
        self.assertTrue(table_block.table_info.has_header_row)
        self.assertIsNotNone(table_block.table_info.cells_raw[0][0].formatting)

    def test_parser_document_contains_style_catalog_snapshot(self) -> None:
        parsed = parse_docx_to_document(input_path=Path("data/input/1.docx"), registry_path=REGISTRY_PATH)

        self.assertGreater(len(parsed.style_catalog), 0)
        style_names = {entry.style_name for entry in parsed.style_catalog}
        self.assertIn("Normal", style_names)

    def test_json_shape_is_machine_readable_with_formatting_fields(self) -> None:
        parsed = parse_docx_to_document(input_path=Path("data/input/1.docx"), registry_path=REGISTRY_PATH)

        payload = json.loads(parsed.model_dump_json())
        self.assertIn("style_catalog", payload)
        self.assertIn("review_render_hints", payload["blocks"][0])
        self.assertIn("runs", payload["blocks"][0])

    @staticmethod
    def _build_paragraph_doc(document: Document) -> None:
        paragraph = document.add_paragraph("Проверка форматирования абзаца")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(9)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(12)

    @staticmethod
    def _build_run_doc(document: Document) -> None:
        paragraph = document.add_paragraph()
        first = paragraph.add_run("Важный")
        first.bold = True
        first.font.name = "Arial"
        first.font.size = Pt(12)
        first.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        second = paragraph.add_run(" текст")
        second.italic = True

    @staticmethod
    def _build_table_doc(document: Document) -> None:
        table = document.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Заголовок"
        table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).text = "Значение"


if __name__ == "__main__":
    unittest.main()
