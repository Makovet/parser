from __future__ import annotations

import unittest

from docx import Document

from qms_doc_parser.parsers.table_parser import parse_table


class TableParserTests(unittest.TestCase):
    def test_parse_plain_table(self) -> None:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = " Header 1 "
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = " A "
        table.cell(1, 1).text = "B"

        parsed = parse_table(table, table_index=1)

        self.assertEqual(parsed.table_id, "tbl_0001")
        self.assertEqual(parsed.rows_count, 2)
        self.assertEqual(parsed.cols_count, 2)
        self.assertEqual(parsed.header_row_count, 1)
        self.assertEqual(parsed.cells_normalized, [["Header 1", "Header 2"], ["A", "B"]])
        self.assertEqual(parsed.cells_raw[0][0].text, "Header 1")
        self.assertEqual(parsed.cells_raw[0][0].row_index, 0)
        self.assertEqual(parsed.cells_raw[0][0].col_index, 0)
        self.assertEqual(parsed.cells_raw[0][0].row_span, 1)
        self.assertEqual(parsed.cells_raw[0][0].col_span, 1)

    def test_parse_table_with_horizontal_merged_cells(self) -> None:
        doc = Document()
        table = doc.add_table(rows=3, cols=3)

        table.cell(0, 0).text = "Merged"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 0).text = "Left"
        table.cell(1, 1).text = "Center"
        table.cell(1, 2).text = "Right"
        table.cell(2, 2).text = "Bottom"

        parsed = parse_table(table, table_index=2)

        self.assertEqual(parsed.rows_count, 3)
        self.assertEqual(parsed.cols_count, 3)
        self.assertEqual(parsed.cells_normalized[0], ["Merged", "Merged", None])

        top_row = parsed.cells_raw[0]
        self.assertEqual(len(top_row), 2)
        self.assertEqual(top_row[0].text, "Merged")
        self.assertEqual(top_row[0].row_index, 0)
        self.assertEqual(top_row[0].col_index, 0)
        self.assertEqual(top_row[0].col_span, 2)
        self.assertEqual(top_row[0].row_span, 1)

    def test_parse_table_with_vertical_merged_cells(self) -> None:
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        table.cell(0, 0).text = "Vertical"
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).text = "H1"
        table.cell(1, 1).text = "H2"
        table.cell(2, 0).text = "BottomL"
        table.cell(2, 1).text = "BottomR"

        parsed = parse_table(table, table_index=3)

        self.assertEqual(parsed.cells_normalized[0], ["Vertical", "H1"])
        self.assertEqual(parsed.cells_normalized[1], ["Vertical", "H2"])

        merged_origin = parsed.cells_raw[0][0]
        self.assertEqual(merged_origin.text, "Vertical")
        self.assertEqual(merged_origin.row_index, 0)
        self.assertEqual(merged_origin.col_index, 0)
        self.assertEqual(merged_origin.row_span, 2)
        self.assertEqual(merged_origin.col_span, 1)

    def test_parse_table_with_mixed_merged_and_empty_cells(self) -> None:
        doc = Document()
        table = doc.add_table(rows=3, cols=3)

        table.cell(0, 0).text = "Top"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 2).text = ""
        table.cell(1, 0).text = ""
        table.cell(1, 1).text = "Mid"
        table.cell(1, 2).text = ""
        table.cell(2, 0).text = "Bottom"
        table.cell(2, 1).text = ""
        table.cell(2, 2).text = "Tail"

        parsed = parse_table(table, table_index=4)

        self.assertEqual(parsed.rows_count, 3)
        self.assertEqual(parsed.cols_count, 3)
        self.assertEqual(parsed.cells_normalized[0], ["Top", "Top", None])
        self.assertEqual(parsed.cells_normalized[1], [None, "Mid", None])
        self.assertEqual(parsed.cells_normalized[2], ["Bottom", None, "Tail"])

        top_origin = parsed.cells_raw[0][0]
        self.assertEqual(top_origin.text, "Top")
        self.assertEqual(top_origin.row_index, 0)
        self.assertEqual(top_origin.col_index, 0)
        self.assertEqual(top_origin.col_span, 2)


if __name__ == "__main__":
    unittest.main()